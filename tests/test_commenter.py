from pathlib import Path

import pytest
from docx import Document

from comment_docx.commenter import (
    CommentDocxError,
    ModelMatch,
    ModelResponse,
    ProviderConfig,
    add_comments_to_docx,
    build_run_context,
    validate_matches,
)


def test_build_run_context_includes_body_and_table_runs_in_order(tmp_path: Path) -> None:
    input_path = tmp_path / "sample.docx"
    document = Document()
    document.add_paragraph("Before")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Inside table"
    document.add_paragraph("After")
    document.save(input_path)

    context = build_run_context(Document(input_path))

    assert [(span.id, span.text) for span in context.spans] == [
        ("r000001", "Before"),
        ("r000002", "Inside table"),
        ("r000003", "After"),
    ]


def test_validate_matches_accepts_start_end_range() -> None:
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("First")
    paragraph.add_run("Second")
    paragraph.add_run("Third")
    context = build_run_context(document)

    selected = validate_matches(
        ModelResponse(
            matches=[
                ModelMatch(
                    start_run_id="r000001",
                    end_run_id="r000003",
                    comment="Explain this.",
                )
            ]
        ),
        context,
    )

    assert [run.text for run in selected[0][0]] == ["First", "Second", "Third"]
    assert selected[0][1] == "Explain this."


def test_validate_matches_accepts_single_run_range() -> None:
    document = Document()
    document.add_paragraph("Only")
    context = build_run_context(document)

    selected = validate_matches(
        ModelResponse(
            matches=[
                ModelMatch(
                    start_run_id="r000001",
                    end_run_id="r000001",
                    comment="Explain this.",
                )
            ]
        ),
        context,
    )

    assert [run.text for run in selected[0][0]] == ["Only"]


def test_validate_matches_rejects_unknown_run_id() -> None:
    document = Document()
    document.add_paragraph("Only")
    context = build_run_context(document)

    with pytest.raises(CommentDocxError, match="unknown run ids"):
        validate_matches(
            ModelResponse(
                matches=[
                    ModelMatch(
                        start_run_id="r000001",
                        end_run_id="r999999",
                        comment="Explain this.",
                    )
                ]
            ),
            context,
        )


def test_validate_matches_rejects_reversed_range() -> None:
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("First")
    paragraph.add_run("Second")
    paragraph.add_run("Third")
    context = build_run_context(document)

    with pytest.raises(CommentDocxError, match="before or equal"):
        validate_matches(
            ModelResponse(
                matches=[
                    ModelMatch(
                        start_run_id="r000003",
                        end_run_id="r000001",
                        comment="Explain this.",
                    )
                ]
            ),
            context,
        )


def test_validate_matches_rejects_blank_comment() -> None:
    document = Document()
    document.add_paragraph("Only")
    context = build_run_context(document)

    with pytest.raises(CommentDocxError, match="empty comment"):
        validate_matches(
            ModelResponse(
                matches=[
                    ModelMatch(
                        start_run_id="r000001",
                        end_run_id="r000001",
                        comment=" ",
                    )
                ]
            ),
            context,
        )


def test_add_comments_to_docx_writes_commented_output(tmp_path: Path) -> None:
    input_path = tmp_path / "sample.docx"
    output_path = tmp_path / "sample.commented.docx"
    document = Document()
    document.add_paragraph("Needs review")
    document.save(input_path)
    provider = ProviderConfig(
        base_url="https://example.com/v1",
        api_key="test-key",
        model="test-model",
    )

    def fake_matcher(context, query: str, provider: ProviderConfig) -> ModelResponse:
        assert query == "Find unclear text"
        assert provider.model == "test-model"
        assert [span.id for span in context.spans] == ["r000001"]
        return ModelResponse(
            matches=[
                ModelMatch(
                    start_run_id="r000001",
                    end_run_id="r000001",
                    comment="Clarify this claim.",
                )
            ]
        )

    result = add_comments_to_docx(
        input_path=input_path,
        output_path=output_path,
        query="Find unclear text",
        provider=provider,
        matcher=fake_matcher,
    )

    commented = Document(output_path)
    comments = list(commented.comments)
    assert result.match_count == 1
    assert result.run_count == 1
    assert comments[0].text == "Clarify this claim."
    assert comments[0].author == "LLM Agent"
