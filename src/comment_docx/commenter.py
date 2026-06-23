from __future__ import annotations

import json
import os
from collections.abc import Callable, Iterable, Iterator
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as DocumentObject
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from openai import BadRequestError, OpenAI, OpenAIError
from pydantic import BaseModel, ConfigDict, Field, ValidationError


class CommentDocxError(Exception):
    """Raised when DOCX commenting cannot be completed safely."""


@dataclass(frozen=True)
class ProviderConfig:
    base_url: str
    api_key: str
    model: str
    timeout: float = 120.0
    max_tokens: int = 2048

    @classmethod
    def from_values(
        cls,
        *,
        base_url: str | None = None,
        api_key: str | None = None,
        model: str | None = None,
        timeout: float = 120.0,
        max_tokens: int = 2048,
    ) -> ProviderConfig:
        resolved_base_url = base_url or os.getenv("PROVIDER_BASE_URL")
        resolved_api_key = api_key or os.getenv("PROVIDER_API_KEY")
        resolved_model = model or os.getenv("PROVIDER_MODEL")

        missing = [
            name
            for name, value in (
                ("PROVIDER_BASE_URL", resolved_base_url),
                ("PROVIDER_API_KEY", resolved_api_key),
                ("PROVIDER_MODEL", resolved_model),
            )
            if not value
        ]
        if missing:
            raise CommentDocxError(
                "Missing OpenAI-compatible provider configuration: " + ", ".join(missing)
            )

        return cls(
            base_url=resolved_base_url,
            api_key=resolved_api_key,
            model=resolved_model,
            timeout=timeout,
            max_tokens=max_tokens,
        )


@dataclass(frozen=True)
class RunSpan:
    id: str
    text: str
    location: str
    order: int
    run: Run

    def to_context_item(self) -> dict[str, str]:
        return {
            "id": self.id,
            "location": self.location,
            "text": self.text,
        }


@dataclass(frozen=True)
class RunContext:
    spans: tuple[RunSpan, ...]

    @property
    def by_id(self) -> dict[str, RunSpan]:
        return {span.id: span for span in self.spans}

    def to_model_payload(self) -> dict[str, Any]:
        return {
            "runs": [span.to_context_item() for span in self.spans],
        }


@dataclass(frozen=True)
class CommentResult:
    input_path: Path
    output_path: Path
    match_count: int
    run_count: int


class ModelMatch(BaseModel):
    model_config = ConfigDict(extra="forbid")

    start_run_id: str = Field(min_length=1)
    end_run_id: str = Field(min_length=1)
    comment: str = Field(min_length=1)


class ModelResponse(BaseModel):
    model_config = ConfigDict(extra="forbid")

    matches: list[ModelMatch] = Field(default_factory=list)


Matcher = Callable[[RunContext, str, ProviderConfig], ModelResponse]


def build_run_context(document: DocumentObject) -> RunContext:
    spans: list[RunSpan] = []

    # AICODE-NOTE: Run ids are the stable contract between the LLM and deterministic DOCX mutation.
    for run in _iter_body_runs(document):
        text = run.text
        if not text.strip():
            continue

        run_id = f"r{len(spans) + 1:06d}"
        spans.append(
            RunSpan(
                id=run_id,
                text=text,
                location=_paragraph_location(run._parent),
                order=len(spans),
                run=run,
            )
        )

    return RunContext(spans=tuple(spans))


def add_comments_to_docx(
    *,
    input_path: Path,
    output_path: Path,
    query: str,
    provider: ProviderConfig,
    author: str = "AI Commenter",
    matcher: Matcher | None = None,
) -> CommentResult:
    document = Document(input_path)
    context = build_run_context(document)
    if not context.spans:
        raise CommentDocxError("DOCX contains no non-empty body or table runs to comment.")

    match_response = (matcher or request_comment_matches)(context, query, provider)
    selected_matches = validate_matches(match_response, context)

    for selected_runs, comment in selected_matches:
        document.add_comment(runs=selected_runs, text=comment, author=author)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)

    return CommentResult(
        input_path=input_path,
        output_path=output_path,
        match_count=len(selected_matches),
        run_count=len(context.spans),
    )


def request_comment_matches(context: RunContext, query: str, provider: ProviderConfig) -> ModelResponse:
    client = OpenAI(api_key=provider.api_key, base_url=provider.base_url, timeout=provider.timeout)
    messages = _build_messages(context=context, query=query)

    try:
        response = client.chat.completions.create(
            model=provider.model,
            messages=messages,
            max_tokens=provider.max_tokens,
            response_format={"type": "json_object"},
        )
    except BadRequestError:
        try:
            response = client.chat.completions.create(
                model=provider.model,
                messages=messages,
                max_tokens=provider.max_tokens,
            )
        except OpenAIError as error:
            raise CommentDocxError(f"LLM request failed: {error}") from error
    except OpenAIError as error:
        raise CommentDocxError(f"LLM request failed: {error}") from error

    content = response.choices[0].message.content
    if not content:
        raise CommentDocxError("LLM response did not include message content.")

    try:
        payload = json.loads(content)
    except json.JSONDecodeError as error:
        raise CommentDocxError("LLM response was not valid JSON.") from error

    try:
        return ModelResponse.model_validate(payload)
    except ValidationError as error:
        raise CommentDocxError(f"LLM response did not match the expected schema: {error}") from error


def validate_matches(response: ModelResponse, context: RunContext) -> list[tuple[list[Run], str]]:
    spans_by_id = context.by_id
    ordered_spans = list(context.spans)
    selected: list[tuple[list[Run], str]] = []

    # AICODE-NOTE: This validation is the safety boundary; the LLM chooses endpoints, code expands ranges.
    for match_index, match in enumerate(response.matches, start=1):
        comment = match.comment.strip()
        if not comment:
            raise CommentDocxError(f"Match {match_index} has an empty comment.")

        endpoint_ids = [match.start_run_id, match.end_run_id]
        unknown_ids = [run_id for run_id in endpoint_ids if run_id not in spans_by_id]
        if unknown_ids:
            raise CommentDocxError(f"Match {match_index} references unknown run ids: {', '.join(unknown_ids)}")

        start_order = spans_by_id[match.start_run_id].order
        end_order = spans_by_id[match.end_run_id].order
        if start_order > end_order:
            raise CommentDocxError(
                f"Match {match_index} start_run_id must be before or equal to end_run_id in document order."
            )

        selected_spans = ordered_spans[start_order : end_order + 1]
        selected.append(([span.run for span in selected_spans], comment))

    return selected


def dump_context(context: RunContext, output_path: Path, query: str) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(
        json.dumps(
            {
                "query": query,
                **context.to_model_payload(),
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )


def default_output_path(input_path: Path) -> Path:
    return input_path.with_name(f"{input_path.stem}.commented{input_path.suffix}")


def _build_messages(*, context: RunContext, query: str) -> list[dict[str, str]]:
    payload = {
        "query": query,
        **context.to_model_payload(),
    }
    return [
        {
            "role": "system",
            "content": (
                "You are a helpful assistant that adds comments on DOCX documents. "
                "The task is to select run id ranges and provide comment texts based on user query. "
                "If the query states that something is missing in the document, "
                "select the appropriate header or paragraph where it supposed to be. "
                "Return only JSON with shape "
                "{\"matches\":[{\"start_run_id\":\"r000001\",\"end_run_id\":\"r000003\",\"comment\":\"...\"}]}. "
                "Use the same id for start_run_id and end_run_id when commenting a single run. "
                "The start_run_id must be before or equal to end_run_id in document order. "
                "If no comment is appropriate, return {\"matches\":[]}."
            ),
        },
        {
            "role": "user",
            "content": json.dumps(payload, ensure_ascii=False),
        },
    ]


def _iter_body_runs(document: DocumentObject) -> Iterator[Run]:
    for block in _iter_inner_content(document):
        if isinstance(block, Paragraph):
            yield from block.runs
        elif isinstance(block, Table):
            yield from _iter_table_runs(block)


def _iter_table_runs(table: Table) -> Iterator[Run]:
    for row_index, row in enumerate(table.rows, start=1):
        for cell_index, cell in enumerate(row.cells, start=1):
            yield from _iter_cell_runs(cell, row_index=row_index, cell_index=cell_index)


def _iter_cell_runs(cell: _Cell, *, row_index: int, cell_index: int) -> Iterator[Run]:
    _ = (row_index, cell_index)
    for block in _iter_inner_content(cell):
        if isinstance(block, Paragraph):
            yield from block.runs
        elif isinstance(block, Table):
            yield from _iter_table_runs(block)


def _iter_inner_content(container: DocumentObject | _Cell) -> Iterable[Paragraph | Table]:
    iter_inner_content = getattr(container, "iter_inner_content", None)
    if iter_inner_content is not None:
        return iter_inner_content()

    return [*container.paragraphs, *container.tables]


def _paragraph_location(paragraph: Paragraph) -> str:
    parent = paragraph._parent
    if isinstance(parent, _Cell):
        return "table-cell paragraph"
    return "body paragraph"
